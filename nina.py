import mysql.connector
import re
from docx import Document


def replace_fields(doc, context):
    for paragraph in doc.paragraphs:
        for field in re.findall(r"{{(.*?)}}", paragraph.text):
            if field in context:
                paragraph.text = paragraph.text.replace("{{" + field + "}}", str(context[field]))

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for field in re.findall(r"{{(.*?)}}", paragraph.text):
                        if field in context:
                            paragraph.text = paragraph.text.replace("{{" + field + "}}", str(context[field]))

# Подключение к базе данных MySQL
mydb = mysql.connector.connect(
    host="localhost",
    user="root",
    password="",
    database="word"
)

mycursor = mydb.cursor()
doc = Document("1121.docx")


# mycursor.execute("SELECT fio, id FROM student")  # Измените запрос на ваш запрос получения студентов
# students = mycursor.fetchall()

# # Создаем пустой список для данных каждой строки таблицы
# table_data = []

# for student in students:
#     context = {
#         'fio': student[0],  # Предполагается, что имя студента находится в первом столбце
#         'id': student[1]    # Предполагается, что ID студента находится во втором столбце
#     }
#     table_data.append(context)  # Добавляем данные каждого студента в список для таблицы

# # Заполняем таблицу в документе
# table = doc.tables[0]  # Предполагается, что таблица находится на первой странице документа

# for row_data in table_data:
#     row = table.add_row().cells
#     replace_fields(doc, row_data)
#     for cell in row:
#         for paragraph in cell.paragraphs:
#             replace_fields(doc, row_data)


# mydb.close()


# Выполнение запроса к базе данных
mycursor.execute("SELECT student_fio, groupe, years, srok, view_practice, direction, name_place, name, class, name_practice, type_practice, number_order, order_date, code_direction FROM end")
results = mycursor.fetchall()
for result in results: 
    context = {
        'student_fio': result[0],
        'groupe_number': result[1],
        'practice_years': result[2],
        'practice_srok': result[3],
        'view_practice_viewe': result[4],
        'direction_name': result[5],
        'place_practice_name_place': result[6],
        'institute_name': result[7],
        'groupe_class': result[8],
        'practice_name_practice': result[9],
        'type_practice': result[10],
        'number_order': result[11],
        'order_date': result[12],
        'code_direction': result[13]

    }
    
    replace_fields(doc, context)
student_fio_value = result[1]

mycursor.execute("SELECT sum_student, plase_practice, city, boss_org, opop, not_st  FROM vse")
results = mycursor.fetchall()
for result in results: 
    context = {
    'sum_student': result[0],
    'plase_practice': result[1],
    'city': result[2],
    'boss_org': result[3],
    'opop': result[4],
    'not_st': result[5]
}

replace_fields(doc, context)



mycursor.execute("SELECT fio FROM student WHERE number_groupe = %s", (student_fio_value,))
fio_results = mycursor.fetchall()
    
for i, fio_result in enumerate(fio_results, start=1):
    context[f'fio_{i}'] = fio_result[0]

replace_fields(doc, context)
mydb.close()
# Закрываем курсор
mycursor.close()

doc.save("final-шаблон2.docx")
