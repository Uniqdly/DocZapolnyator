import mysql.connector
import re
from docx import Document

def replace_fields(doc, context):
    for paragraph in doc.paragraphs:
        for field in re.findall(r"{{(.*?)}}", paragraph.text):
            if field in context:
                paragraph.text = paragraph.text.replace("{{" + field + "}}", str(context[field]))

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for field in re.findall(r"{{(.*?)}}", paragraph.text):
                        if field in context:
                            paragraph.text = paragraph.text.replace("{{" + field + "}}", str(context[field]))

# Подключение к базе данных MySQL
mydb = mysql.connector.connect(
  host="localhost",
  user="root",
  password="",
  database="word"
)

mycursor = mydb.cursor()
doc = Document("шаблон.docx")

# Выполнение запроса к базе данных
mycursor.execute("SELECT student_fio,hards, quality, size_work, comments, rate, groupe, years, srok, view_practice, direction, address, name_place, name, class, name_practice,Ugra_boss_fio,Ugra_boss_post,company_boss_fio,company_boss_post FROM end")
result = mycursor.fetchone()
results = mycursor.fetchall()

if results:
    result = results[-1]
context = {
            'student_fio': result[0],
            'practice_student_hards': result[1],
            'practice_student_quality': result[2],
            'practice_student_size_work': result[3],
            'comments_text': result[4],
            'practice_student_rate': result[5],
            'groupe_number': result[6],
            'practice_years': result[7],
            'practice_srok': result[8],
            'view_practice_viewe': result[9],
            'direction_name': result[10],
            'place_practice_address': result[11],
            'place_practice_name_place': result[12],
            'institute_name': result[13],
            'groupe_class': result[14],
            'name_practice': result[15],
            'boss_practice_ugrasu_fio': result[16],
            'boss_practice_ugrasu_post': result[17],
            'boss_practice_company_fio': result[18],
            'boss_practice_company_post': result[19],
            'boss_practice_org_company_fio': result[18],
            'boss_practice_org_company_post': result[19]
        }

student_fio_value = result[0]
mycursor.fetchall()
mycursor.execute("SELECT name_task FROM tasks WHERE student_fio = %s", (student_fio_value,))
tasks_results = mycursor.fetchall()
    
for i, task_result in enumerate(tasks_results, start=1):
    context[f'tasks_name_task_{i}'] = task_result[0]


mycursor.fetchall()
mycursor.execute("SELECT datee FROM tasks WHERE student_fio = %s", (student_fio_value,))
datee_results = mycursor.fetchall()
    
for i, datee_result in enumerate(datee_results, start=1):
    context[f'tasks_date_{i}'] = datee_result[0]


replace_fields(doc, context)
doc.save("шаблон-final.docx")